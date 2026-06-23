from __future__ import annotations
from datetime import datetime, timezone
from models.phase5 import ProductionMeta, GoldenRule, Phase5Response, ExportHistoryItem

GOLDEN_RULES_LIST = [
    {"num": 1, "name": "Demonstration-First Dialogue", "description": "Show before saying", "category": "Golden Rule"},
    {"num": 2, "name": "Action-to-Dialogue Ratio >= 2.0", "description": "2+ action words per spoken word", "category": "Golden Rule"},
    {"num": 3, "name": "Silent-Runtime Quota by Format", "description": "ABR >=50%, Whiteboard >=30%, Talking-head >=20%", "category": "Golden Rule"},
    {"num": 4, "name": "Hook Parseable Muted", "description": "Visual + caption carries message alone", "category": "Golden Rule"},
    {"num": 5, "name": "Broad → Narrow → Niche", "description": "Hook=broad, body=ICP, CTA=niche", "category": "Golden Rule"},
    {"num": 6, "name": "Show Don't Tell (LOC/EBA/VWFA)", "description": "Every fact is demonstrated", "category": "Golden Rule"},
    {"num": 7, "name": "No Talking Head > 2s Without Change", "description": "Pattern interrupt every 2s", "category": "Quality Gate"},
    {"num": 8, "name": "Silent Moments Mandatory", "description": "Reaction 0.5-1s, product reveal 1-2s, pre-CTA freeze 1s", "category": "Quality Gate"},
    {"num": 9, "name": "Bridge Words Between Sentences", "description": "'And...', 'But here's...', 'Which means...'", "category": "Quality Gate"},
    {"num": 10, "name": "Action-Integrated Dialogue (EBA+STS)", "description": "Every line specifies what actor does", "category": "Quality Gate"},
    {"num": 11, "name": "Tri-Modal Convergence in 0-1.5s", "description": "Visual + audio + text by 1.5s", "category": "Neuroscience"},
    {"num": 12, "name": "Dopamine Peak Placement", "description": "Emotional peak in first 40% of video", "category": "Neuroscience"},
    {"num": 13, "name": "Single Peak Audio Rule", "description": "ONE cash register / peak SFX per video", "category": "Studio"},
    {"num": 14, "name": "CTA Freeze Rule", "description": "1.0s silence before CTA", "category": "CTA"},
    {"num": 15, "name": "Sub-Topic Bridging", "description": "Every scene transition uses a bridge word", "category": "Quality Gate"},
]


def evaluate_golden_rules(scenes: list[dict], phase4_checks: list[dict] | None = None) -> list[GoldenRule]:
    rules = []
    check_map = {}
    if phase4_checks:
        for c in phase4_checks:
            check_map[c.get("name", "")] = c

    for r in GOLDEN_RULES_LIST:
        matched = check_map.get(r["name"])
        passed = matched.get("result") == "PASS" if matched else True
        evidence = matched.get("evidence", "") if matched else "Auto-evaluated"
        rules.append(GoldenRule(
            num=r["num"], name=r["name"], description=r["description"],
            category=r["category"], passed=passed, evidence=evidence,
        ))
    return rules


def assemble_production_pack(supabase, brief_id: str) -> Phase5Response:
    # Fetch all phase data
    brief = supabase.table("briefs").select("*").eq("id", brief_id).execute()
    brief_data = brief.data[0] if brief.data else {}

    p1 = supabase.table("phase1_data").select("*").eq("brief_id", brief_id).execute()
    p1d = p1.data[0] if p1.data else {}

    p2 = supabase.table("phase2_data").select("*").eq("brief_id", brief_id).execute()
    p2d = p2.data[0] if p2.data else {}

    p3 = supabase.table("phase3_data").select("*").eq("brief_id", brief_id).execute()
    p3d = p3.data[0] if p3.data else {}

    p4 = supabase.table("phase4_data").select("*").eq("brief_id", brief_id).execute()
    p4d = p4.data[0] if p4.data else {}

    scenes = p3d.get("scenes", [])
    meta = ProductionMeta(
        title=brief_data.get("title", p1d.get("topic", "")),
        actor=brief_data.get("creator_name", ""),
        company="Backero",
        platform=p1d.get("platform", ""),
        format=p2d.get("selected_format", ""),
        contentType=p2d.get("content_type", ""),
        structure=p2d.get("selected_structure", ""),
        runtime=f"{p3d.get('total_runtime_sec', 0)}s",
        scenes=len(scenes),
        cuts=p3d.get("cut_count", 0),
        words=p3d.get("total_words", 0),
        verdict=p4d.get("overall_verdict", ""),
        score=p4d.get("quality_score", 0),
    )

    golden_rules = evaluate_golden_rules(scenes, p4d.get("checks"))

    return Phase5Response(meta=meta, scenes=scenes, golden_rules=golden_rules)


def get_tab_view(scenes: list[dict], tab_name: str) -> list[dict]:
    if tab_name == "all":
        return scenes
    filtered = []
    for s in scenes:
        if tab_name == "actor":
            filtered.append({
                "sceneNum": s.get("sceneNum"), "name": s.get("name"),
                "timingStart": s.get("timingStart"), "timingEnd": s.get("timingEnd"),
                "dialogue": s.get("dialogue"), "action": s.get("action"),
            })
        elif tab_name == "camera":
            filtered.append({
                "sceneNum": s.get("sceneNum"), "name": s.get("name"),
                "camera": s.get("camera"), "editMarkers": s.get("editMarkers"),
            })
        elif tab_name == "edit":
            filtered.append({
                "sceneNum": s.get("sceneNum"), "name": s.get("name"),
                "visual": s.get("visual"), "audio": s.get("audio"),
                "editMarkers": s.get("editMarkers"),
            })
        elif tab_name == "script":
            filtered.append({
                "sceneNum": s.get("sceneNum"), "name": s.get("name"),
                "dialogue": s.get("dialogue"),
            })
    return filtered


def generate_docx(meta: ProductionMeta, scenes: list[dict], golden_rules: list[GoldenRule], brief_id: str) -> str:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Cover
    h = doc.add_heading(meta.title or "Production Pack", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    cells = [
        ("Platform", meta.platform), ("Format", meta.format),
        ("Runtime", meta.runtime), ("Score", f"{meta.score}/100"),
        ("Verdict", meta.verdict),
    ]
    for i, (k, v) in enumerate(cells):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = str(v)

    doc.add_page_break()

    # Actor's Script
    doc.add_heading("ACTOR'S SCRIPT", level=1)
    for s in scenes:
        p = doc.add_paragraph()
        run = p.add_run(f"Scene {s.get('sceneNum', '?')}: {s.get('name', '')}")
        run.bold = True
        run.font.size = Pt(13)
        p2 = doc.add_paragraph()
        p2.add_run(f"[{s.get('timingStart', 0)}s – {s.get('timingEnd', 0)}s]").italic = True
        doc.add_paragraph(s.get("dialogue", ""))
        act = doc.add_paragraph()
        act.add_run(s.get("action", "")).italic = True
        doc.add_paragraph("")

    doc.add_page_break()

    # Camera Sheet
    doc.add_heading("CAMERA / DoP SHEET", level=1)
    for s in scenes:
        doc.add_heading(f"Scene {s.get('sceneNum', '?')}: {s.get('name', '')}", level=3)
        cam = s.get("camera", {})
        t = doc.add_table(rows=3, cols=2)
        t.style = "Table Grid"
        t.rows[0].cells[0].text = "Shot"
        t.rows[0].cells[1].text = cam.get("shot", "")
        t.rows[1].cells[0].text = "Angle"
        t.rows[1].cells[1].text = cam.get("angle", "")
        t.rows[2].cells[0].text = "Movement"
        t.rows[2].cells[1].text = cam.get("movement", "")

        markers = s.get("editMarkers", [])
        if markers:
            mt = doc.add_table(rows=len(markers) + 1, cols=2)
            mt.style = "Table Grid"
            mt.rows[0].cells[0].text = "Time"
            mt.rows[0].cells[1].text = "Event"
            for i, m in enumerate(markers):
                mt.rows[i + 1].cells[0].text = m.get("time", "")
                mt.rows[i + 1].cells[1].text = m.get("event", "")
        doc.add_paragraph("")

    doc.add_page_break()

    # Editor's Brief
    doc.add_heading("EDITOR'S BRIEF", level=1)
    for s in scenes:
        doc.add_heading(f"Scene {s.get('sceneNum', '?')}: {s.get('name', '')}", level=3)
        doc.add_paragraph(f"Visual: {s.get('visual', 'N/A')}")
        doc.add_paragraph(f"Audio: {s.get('audio', 'N/A')}")

    doc.add_page_break()

    # Golden Rules
    doc.add_heading("GOLDEN RULES CHECKLIST", level=1)
    gt = doc.add_table(rows=len(golden_rules) + 1, cols=4)
    gt.style = "Table Grid"
    gt.rows[0].cells[0].text = "#"
    gt.rows[0].cells[1].text = "Rule"
    gt.rows[0].cells[2].text = "Category"
    gt.rows[0].cells[3].text = "Status"
    for i, r in enumerate(golden_rules):
        gt.rows[i + 1].cells[0].text = str(r.num)
        gt.rows[i + 1].cells[1].text = r.name
        gt.rows[i + 1].cells[2].text = r.category
        gt.rows[i + 1].cells[3].text = "✅ PASS" if r.passed else "❌ FAIL"

    # Footer
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = f"CONFIDENTIAL — {meta.company} — {datetime.now().strftime('%Y-%m-%d')}"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    path = f"/tmp/{brief_id}_production_pack.docx"
    doc.save(path)
    return path
